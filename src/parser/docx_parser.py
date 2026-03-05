"""Парсинг .docx протоколов для извлечения числа участников.

Стратегии (в порядке приоритета):
1. Поиск в тексте параграфов через regex-паттерны
2. Анализ таблиц — извлечение номеров заявок и дедупликация
3. Подсчёт строк в таблицах с заголовками «Участник»/«Заявка»/«Поставщик»
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from loguru import logger

from src.parser.participant_patterns import (
    ALL_PARTICIPANT_HEADERS,
    ParticipantParsingResult,
    ParticipantResult,
    TABLE_HEADERS,
    extract_participants_from_text,
)


# ── Дата-классы ───────────────────────────────────────────────────────────────


@dataclass
class TenderApplications:
    """Хранит уникальные номера заявок для одного тендера.

    Используется для дедупликации заявок между несколькими таблицами
    одного протокола или между несколькими протоколами одного тендера.
    """

    tender_id: str
    applications: set[int] = field(default_factory=set)

    def add_from_table(self, numbers: list[int]) -> None:
        """Добавляет номера заявок из таблицы (дедупликация автоматическая)."""
        self.applications.update(numbers)

    def get_unique_count(self) -> int:
        """Возвращает количество уникальных заявок."""
        return len(self.applications)


# ── Вспомогательные функции ───────────────────────────────────────────────────


def find_table_title_before(doc, table_idx: int) -> str | None:
    """Находит текст (заголовок) непосредственно перед таблицей в DOCX.

    DOCX документ хранит параграфы и таблицы как последовательные элементы
    в body. Эта функция ищет параграф, стоящий непосредственно перед
    таблицей с индексом ``table_idx``.

    Args:
        doc: python-docx Document объект.
        table_idx: Индекс таблицы в doc.tables.

    Returns:
        Текст параграфа-заголовка или None, если не найден.
    """
    try:
        from docx.oxml.ns import qn  # type: ignore[import-untyped]
    except ImportError:
        return None

    body = doc.element.body
    # Собираем все дочерние элементы body
    children = list(body)

    # Находим XML-элемент нашей таблицы
    table_elements = [child for child in children if child.tag == qn("w:tbl")]

    if table_idx >= len(table_elements):
        return None

    target_table_xml = table_elements[table_idx]
    target_pos = list(children).index(target_table_xml)

    # Ищем предыдущие параграфы (максимум 3 шага назад)
    for i in range(target_pos - 1, max(-1, target_pos - 4), -1):
        child = children[i]
        if child.tag == qn("w:p"):
            # Извлекаем текст параграфа
            para_text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if para_text:
                return para_text
        elif child.tag == qn("w:tbl"):
            # Предыдущий элемент — другая таблица, останавливаемся
            break

    return None


def extract_application_numbers_from_table(table) -> list[int]:
    """Извлекает номера заявок из таблицы DOCX.

    Предполагает, что первый столбец таблицы содержит порядковые номера
    заявок (1, 2, 3, ...). Пропускает строку заголовка.

    Args:
        table: python-docx Table объект.

    Returns:
        Отсортированный список номеров заявок.
    """
    numbers: list[int] = []
    if not table.rows or len(table.rows) < 2:
        return numbers

    for row in table.rows[1:]:  # Пропустить заголовок
        if not row.cells:
            continue
        first_cell = row.cells[0].text.strip()
        match = re.match(r"^(\d+)", first_cell)
        if match:
            num = int(match.group(1))
            numbers.append(num)

    return sorted(set(numbers))  # Дедупликация и сортировка


def _is_participant_table(table, doc=None, table_idx: int = -1) -> bool:
    """Проверяет, является ли таблица таблицей с заявками/участниками.

    Проверяет:
    1. Заголовки ячеек первой строки таблицы
    2. Текст перед таблицей (если передан doc и table_idx)

    Args:
        table: python-docx Table объект.
        doc: python-docx Document (необязательно, для поиска заголовка перед таблицей).
        table_idx: Индекс таблицы в doc.tables (необязательно).

    Returns:
        True если таблица содержит заявки/участников.
    """
    if not table.rows:
        return False

    # 1. Проверяем заголовок первой строки таблицы
    header_row = table.rows[0]
    header_cells = [cell.text.strip().lower() for cell in header_row.cells]

    has_participant_header = any(
        any(ph in cell_text for ph in ALL_PARTICIPANT_HEADERS)
        for cell_text in header_cells
    )

    if has_participant_header:
        return True

    # 2. Проверяем текст над таблицей
    if doc is not None and table_idx >= 0:
        title_before = find_table_title_before(doc, table_idx)
        if title_before:
            title_lower = title_before.lower()
            if any(ph in title_lower for ph in ALL_PARTICIPANT_HEADERS):
                logger.debug(
                    "Таблица #{} идентифицирована по заголовку перед ней: '{}'",
                    table_idx,
                    title_before[:60],
                )
                return True

    return False


# ── Основные функции ──────────────────────────────────────────────────────────


def extract_participants_from_docx(file_path: Path) -> ParticipantParsingResult:
    """Извлекает количество участников из .docx файла протокола.

    Args:
        file_path: Путь к .docx файлу.

    Returns:
        ParticipantParsingResult с количеством и методом извлечения.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        logger.error("python-docx не установлен, невозможно парсить .docx")
        return ParticipantParsingResult(
            count=None, numbers=[], method="docx_import_error", confidence="low"
        )

    try:
        doc = Document(str(file_path))
    except Exception as exc:
        logger.error("Ошибка при открытии .docx файла {}: {}", file_path, exc)
        return ParticipantParsingResult(
            count=None, numbers=[], method="docx_open_error", confidence="low"
        )

    # ── Стратегия 1: Собираем весь текст параграфов + текст таблиц ───────
    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text_parts.append(text)

    # Текст из таблиц тоже добавляем (часто участники перечислены в таблицах)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                all_text_parts.append(" | ".join(row_texts))

    full_text = "\n".join(all_text_parts)

    if not full_text.strip():
        logger.warning("Файл {} не содержит текста", file_path.name)
        return ParticipantParsingResult(
            count=None, numbers=[], method="docx_empty", confidence="low"
        )

    logger.debug(
        "DOCX {}: {} параграфов, {} таблиц, {} символов текста",
        file_path.name,
        len(doc.paragraphs),
        len(doc.tables),
        len(full_text),
    )

    # ── Стратегия 1: regex-паттерны по тексту ────────────────────────────
    result = extract_participants_from_text(full_text)
    if result.count is not None:
        result.method = f"docx_{result.method}"
        return result

    # ── Стратегия 2: Анализ таблиц с дедупликацией по номерам заявок ─────
    table_result = _analyze_tables(doc)
    if table_result is not None:
        return table_result

    # Ничего не найдено
    logger.info("DOCX {}: не удалось определить число участников", file_path.name)
    return ParticipantParsingResult(
        count=None, numbers=[], method="docx_no_pattern", confidence="low"
    )


def _analyze_tables(doc) -> ParticipantParsingResult | None:
    """Анализирует таблицы .docx документа для подсчёта участников.

    Ищет таблицы с заголовками, содержащими «Участник», «Заявка»,
    «Поставщик» и т.п. Извлекает номера заявок из первого столбца
    и выполняет дедупликацию между несколькими таблицами.

    Args:
        doc: python-docx Document объект.

    Returns:
        ParticipantParsingResult или None если таблиц с участниками не найдено.
    """
    all_application_numbers: set[int] = set()
    total_data_rows = 0
    tables_found = 0

    for table_idx, table in enumerate(doc.tables):
        if not _is_participant_table(table, doc=doc, table_idx=table_idx):
            continue

        tables_found += 1

        # Извлекаем номера заявок из первого столбца
        app_numbers = extract_application_numbers_from_table(table)

        if app_numbers:
            # Дедупликация между таблицами
            new_numbers = set(app_numbers) - all_application_numbers
            all_application_numbers.update(app_numbers)
            logger.debug(
                "Таблица #{}: номера заявок {}, новых: {}",
                table_idx,
                app_numbers,
                sorted(new_numbers),
            )
        else:
            # Нет номеров — считаем строки данных
            data_rows = 0
            for row in table.rows[1:]:
                row_text = " ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip() and not row_text.lower().startswith("итого"):
                    data_rows += 1

            if data_rows > 0:
                total_data_rows += data_rows
                logger.debug(
                    "Таблица #{}: без номеров заявок, строк данных: {}",
                    table_idx,
                    data_rows,
                )

    if tables_found == 0:
        return None

    # Определяем итоговое количество
    if all_application_numbers:
        unique_count = len(all_application_numbers)
        sorted_numbers = sorted(all_application_numbers)
        logger.debug(
            "Итого: {} уникальных заявок из {} таблиц(ы): {}",
            unique_count,
            tables_found,
            sorted_numbers,
        )
        return ParticipantParsingResult(
            count=unique_count,
            numbers=sorted_numbers,
            method="docx_table_participant_rows",
            confidence="medium",
        )
    elif total_data_rows > 0:
        logger.debug(
            "Итого: {} строк данных из {} таблиц(ы) (без номеров заявок)",
            total_data_rows,
            tables_found,
        )
        return ParticipantParsingResult(
            count=total_data_rows,
            numbers=[],
            method="docx_table_participant_rows",
            confidence="medium",
        )

    return None
